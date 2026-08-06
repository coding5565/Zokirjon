"""
Генерация .docx-файла с фидбеком по Writing/Speaking.

По просьбе заказчика фидбек отправляется файлом, а не длинной стеной текста
в чате — аудитория (Gen-Z студенты) такие сообщения не дочитывает, а файл
воспринимается как «настоящий» разбор, который можно спокойно открыть и
прочитать.

Модель форматирует свой ответ лёгким markdown'ом (### заголовки, **bold**,
- списки, | таблицы |) — раньше build_feedback_document() просто резал текст
по пустым строкам и клал как есть, из-за чего в готовом файле были видны
сырые "##"/"**" символы. Заказчик явно попросил это убрать ("AI ekanligini
bildirib turibdi" — по этим символам видно, что текст сгенерирован ИИ) и
использовать вместо них настоящее форматирование Word: жирный текст без
звёздочек для заголовков/меток, реальный маркированный список (чёрная точка)
для пунктов списка, реальную таблицу для markdown-таблиц. Плюс межстрочный
интервал 1.5 по всему документу (тоже прямая просьба заказчика).

Это лёгкий построчный парсер, не полноценный markdown-движок — ровно
настолько, насколько модель реально форматирует свой вывод (см.
modules/ai_feedback.py, все *_STRUCTURE-константы).
"""

import io
import re

from docx import Document

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_HEADER_RE = re.compile(r"^#{1,6}\s+(.*)")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^[\s|:-]+$")

# Заказчик прямо попросил интервал 1.5 по всему документу.
_LINE_SPACING = 1.5


def _add_inline_runs(paragraph, text: str, force_bold: bool = False) -> None:
    """
    Разбивает строку по **bold**-маркерам на runs — сами звёздочки в готовый
    текст не попадают, только настоящее bold-форматирование Word.
    """
    parts = _BOLD_RE.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = force_bold or (i % 2 == 1)


def _new_paragraph(document: Document, style: str = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = _LINE_SPACING
    return paragraph


def _add_table(document: Document, table_lines: list[str]) -> None:
    """markdown-таблица (| a | b |) -> настоящая таблица Word, не сырой текст с '|'."""
    rows = []
    for line in table_lines:
        if _TABLE_SEP_RE.match(line.strip("|")):
            continue  # строка-разделитель |---|---|
        rows.append([cell.strip() for cell in line.strip("|").split("|")])
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    for row_cells in rows:
        row = table.add_row()
        for col_index in range(col_count):
            text = row_cells[col_index] if col_index < len(row_cells) else ""
            _add_inline_runs(row.cells[col_index].paragraphs[0], text)


def build_feedback_document(feedback_text: str, title: str) -> bytes:
    document = Document()
    document.styles["Normal"].paragraph_format.line_spacing = _LINE_SPACING
    document.add_heading(title, level=1).paragraph_format.line_spacing = _LINE_SPACING

    lines = feedback_text.split("\n")
    i, total = 0, len(lines)
    while i < total:
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if _TABLE_ROW_RE.match(line):
            table_lines = []
            while i < total and _TABLE_ROW_RE.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(document, table_lines)
            continue

        header_match = _HEADER_RE.match(line)
        if header_match:
            _add_inline_runs(_new_paragraph(document), header_match.group(1), force_bold=True)
            i += 1
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            _add_inline_runs(_new_paragraph(document, style="List Bullet"), bullet_match.group(1))
            i += 1
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            _add_inline_runs(_new_paragraph(document, style="List Number"), numbered_match.group(1))
            i += 1
            continue

        _add_inline_runs(_new_paragraph(document), line)
        i += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
